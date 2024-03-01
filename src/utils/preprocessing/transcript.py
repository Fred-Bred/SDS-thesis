"""
Utility functions for working with transcripts.
"""
import os
import random
import zipfile
import re

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.opc.exceptions import PackageNotFoundError
import xml.etree.ElementTree as ET
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

import pandas as pd
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from faker import Faker

# Function to extract text from a docx file
def get_docx_text(filename, class_df=False):
    """Extract text from a docx file"""
    document = Document(filename)

    docText = '\n\n'.join([paragraph.text for paragraph in document.paragraphs])

    if class_df:
        doc_class = filename.split('_')[-1].split('.')[0]
        return pd.DataFrame({'text': [docText], 'class': [doc_class]})
    else:
        return docText

# Function to extract comments from a docx file
def extract_comments(filename):
    """Extract all comments from a docx file"""
    doc = Document(filename)
    try:
        comments_part = doc.part.package.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
    except KeyError:
        return []
    root = ET.fromstring(comments_part.blob)
    comments = [elem.text for elem in root.iter() if elem.tag.endswith('comment')]
    return comments

# Function to extract all the comments of document 
def get_document_comments(docxFileName):
    """Returns a dictionary with comment id as key and comment string as value"""
    ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    commentsXML = docxZip.read('word/comments.xml')
    et = etree.XML(commentsXML)
    comments = et.xpath('//w:comment',namespaces=ooXMLns)
    for c in comments:
        comment=c.xpath('string(.)',namespaces=ooXMLns)
        comment_id=c.xpath('@w:id',namespaces=ooXMLns)[0]
        comments_dict[comment_id]=comment
    return comments_dict

#Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph, comments_dict):
    """Returns a list of comments in a given paragraph"""
    comments=[]
    for run in paragraph.runs:
        comment_reference=run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id=comment_reference[0].xpath('@w:id')[0]
            comment=comments_dict[comment_id]
            comments.append(comment)
    return comments

# Function to fetch all paragraphs with comments from a document
def comments_with_reference_paragraph(docxFileName):
    """Returns a dict with keys as paragraphs and values as comments"""
    document = Document(docxFileName)
    comments_dict=get_document_comments(docxFileName)
    comments_with_their_reference_paragraph={}
    for paragraph in document.paragraphs:  
        if comments_dict: 
            comments=paragraph_comments(paragraph,comments_dict)  
            if comments:
                comments_with_their_reference_paragraph[paragraph.text] = comments
    return comments_with_their_reference_paragraph

# Function to extract lists of text and comments from a dict where keys are paragraphs and values are comments
def extract_text_and_comments(comments_with_reference_paragraph, as_lists=False):
    """Returns a list or df of text and comments from a list of dicts where keys are paragraphs and values are comments"""
    text_list = []
    comments_list = []
    for item in comments_with_reference_paragraph:
        for key, value in item.items():
            text_list.append(key)
            comments_list.append(value)
    # text_list = [item for sublist in text_list for item in sublist]
    comments_list = [item for sublist in comments_list for item in sublist]
    # Split label strings by comma
    comments_list = [i.split(", ") for i in comments_list]
    if as_lists:
        return text_list, comments_list
    else:
        # Create df
        df = pd.DataFrame(list(zip(text_list, comments_list)), columns = ["sentences", "labels"])
        
        # Convert the list of labels into a series of strings (so that get_dummies can process it)
        df['labels'] = df['labels'].apply(lambda x: ','.join(map(str, x)))

        # Convert the labels column into dummy variables
        df = df.join(df['labels'].str.get_dummies(','))

        # Drop the labels column
        df = df.drop(columns=['labels'])

        return df

# Load data from folder
def load_data_from_folder(folder_path, as_lists=False, subscales=False):
    """Load data from folder
    args:
        folder_path (str): path to folder containing docx files
        as_lists (bool): return lists instead of df
        subscales (bool): return subscales df instead of full df
    """
    # Get all the docx files from the folder
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

    # Extract text and comments from each docx file
    docs = [comments_with_reference_paragraph(f'{folder_path}/{f}') for f in docx_files]

    # Extract text and comments
    text_list, comments_list = extract_text_and_comments(docs, as_lists=True)

    if as_lists:
        return text_list, comments_list
    else:
        # Create df
        df = pd.DataFrame(list(zip(text_list, comments_list)), columns = ["sentence", "labels"])
        
        # Convert the list of labels into a series of strings (so that get_dummies can process it)
        df['labels'] = df['labels'].apply(lambda x: ','.join(map(str, x)))

        # Remove leading 0 from labels with regex
        df['labels'] = df['labels'].str.replace(r'\b0+', '', regex=True)

        # Convert the labels column into dummy variables
        df = df.join(df['labels'].str.get_dummies(','))

        # Drop the labels column
        df = df.drop(columns=['labels'])

        # Get a list of the column names (excluding the first one) and sort it
        columns = df.columns[1:].to_list()
        columns.sort(key=int)

        # Index the DataFrame with the first column name and the sorted list of other column names
        df = df[[df.columns[0]] + columns]

        if subscales: # create subscale df
            # Define the scale mapping
            scale_mapping = {
                1: list(range(1, 11)),
                2: list(range(11, 16)),
                3: list(range(16, 25)),
                4: list(range(25, 39)),
                5: list(range(39, 60))
            }

            # Initialize the new DataFrame with zeros
            scale_df = pd.DataFrame(data=np.zeros((len(df), len(scale_mapping))), columns=[f"scale_{i}" for i in range(1, len(scale_mapping)+1)])

            # Add the 'sentence' column from df to scale_df
            scale_df.insert(0, 'sentence', df['sentence'])

            # Iterate over the scale mapping
            for scale, scale_nums in scale_mapping.items():
                # Get the columns in df that belong to this scale
                scale_cols = [col for col in df.columns[1:] if int(col) in scale_nums]
                # If any of these columns have a value in a row, set the corresponding scale in scale_df to 1
                scale_df[f'scale_{scale}'] = df[scale_cols].any(axis=1).astype(int)
            return scale_df
        else:
            return df


# Function to create fake data
# Function to create a random Word document
def create_random_word_document(folder_path, n_docs=1):
    """Create n random Word documents"""
    fake = Faker()
    letters = ['A', 'B', 'C']

    for i in range(n_docs):
        doc = Document()
        filename = f'{folder_path}/random_doc_{i}_{random.choice(letters)}.docx'

        # Add 3-6 paragraphs
        for _ in range(random.randint(3, 6)):
            paragraph = doc.add_paragraph(fake.text())
            paragraph.add_comment(str(random.randint(0, 59)), author=fake.name(), initials='ab')

        doc.save(filename)

# Function to load patient speech turns from path or doc
def load_patient_turns(doc):
    """Load the patient speech turns from a document.
    Args:
        doc (str or Document): The document to load.
        prefix (str): The prefix used to indicate a patient speech turn."""

    # Load the document if doc is a path
    if isinstance(doc, str):
        if doc.endswith('.docx'):
            doc = Document(doc)
            # Extract the text of each paragraph
            paragraphs = [p.text for p in doc.paragraphs]
        else:
            raise ValueError("Unsupported file type. Please provide a .docx file.")

        # Compile the regular expression for matching the prefixes
        prefix_re = re.compile(r'^(P\d*:|P:|PATIENT:|P;|PATIENT;)')
        
        # Filter the paragraphs to only include those that start with the prefix
        paragraphs = [p for p in paragraphs if prefix_re.match(p)]

        # Strip the prefix from the paragraphs
        paragraphs = [prefix_re.sub('', p) for p in paragraphs]

    return paragraphs

# Function to load patient speech turns from all documents in a folder
def load_patient_turns_from_folder(folder_path, prefixes=['P:', 'PATIENT:', 'P;', 'PATIENT;']):
    """Load the patient speech turns from all documents in a folder.
    Returns a list of lists with speech turns per document.
    Args:
        folder_path (str): The path to the folder containing the documents.
        prefixes (list): The prefixes used to indicate a patient speech turn."""

    # Initialize an empty list to hold all the paragraphs
    all_paragraphs = []

    # Compile the regular expression for matching the prefixes
    prefix_re = re.compile(r'^(P\d*:|P:|PATIENT:|P;|PATIENT;)')

    # Iterate over all files in the folder
    for filename in os.listdir(folder_path):
        # Check if the file is a Word document
        if filename.endswith('.docx'):
            # Load the document
            doc = Document(os.path.join(folder_path, filename))

            # Extract the text of each paragraph
            paragraphs = [p.text for p in doc.paragraphs]

            # Filter the paragraphs to only include those that start with the prefix
            paragraphs = [p for p in paragraphs if prefix_re.match(p)]

            # Strip the prefix from the paragraphs
            paragraphs = [prefix_re.sub('', p) for p in paragraphs]

            # Add the paragraphs to the list of all paragraphs
            all_paragraphs.append(paragraphs)

    return all_paragraphs

# Function to load patient speech turns from all documents in a folder and store in df with labels
def load_data_with_labels(labels_path, folder_path):
    """Load the patient speech turns from all documents in a folder and store them in a DataFrame with labels.
    Args:
        labels_path (str): The path to the Excel sheet with the labels.
        folder_path (str): The path to the folder containing the documents.
    Returns:
        A DataFrame with the patient speech turns and their labels."""
    # Load the Excel sheet
    df = pd.read_excel(labels_path)

    # Counter for the number of documents check
    n_docs = 0

    # Initialize a list to store DataFrame objects
    df_list = []

    # Iterate over the rows in the DataFrame
    for index, row in df.iterrows():
        # Get the document name and label
        document_name = row['Document']
        label = row['Class3']

        # Load the corresponding file
        try:
            p_turns = load_patient_turns(os.path.join(folder_path, document_name))

            # Increment counter
            n_docs += 1
        except FileNotFoundError:
            continue
        except PackageNotFoundError:
            continue

        # For each patient turn, create a new row in the result DataFrame
        data = []
        for patient_turn in p_turns:
            data.append({'text': patient_turn, 'label': label})

        # Append the DataFrame to the list
        df_list.append(pd.DataFrame(data))
    
    # Concatenate all the DataFrames in the list
    result = pd.concat(df_list, ignore_index=True)

    print(f"\nLoaded {n_docs} documents.")
    return result

# Get average word count of speech turns
def average_word_count(input_list):
    """Calculate the average word count of the strings in a list of lists of strings or a list of strings.
    Args:
        input_list (list): A list of lists of strings or a list of strings.
    Returns:
        The average word count of the strings."""

    # Initialize counters for the total word count and the total number of strings
    total_word_count = 0
    total_strings = 0

    # Check if the input is a list of lists
    if isinstance(input_list[0], list):
        # Iterate over all lists in the list of lists
        for list_ in input_list:
            # Iterate over all strings in the list
            for string in list_:
                # Increment the total word count by the word count of the string
                total_word_count += len(string.split())

                # Increment the total number of strings
                total_strings += 1
    else:
        # The input is a list of strings
        # Iterate over all strings in the list
        for string in input_list:
            # Increment the total word count by the word count of the string
            total_word_count += len(string.split())

            # Increment the total number of strings
            total_strings += 1

    # Calculate the average word count
    average_word_count = total_word_count / total_strings

    return average_word_count

# Function to discard speech turns under a certain length
def filter_by_word_count(data, min_word_count=100):
    """Filter out strings that are under a specified word count.
    Args:
        data (list): A list of strings or a list of lists of strings.
        min_word_count (int): The minimum word count.
    Returns:
        A list of strings or a list of lists of strings, without strings that are under the specified word count."""

    if all(isinstance(i, str) for i in data):
        # If data is a list of strings, filter out strings that are under the specified word count
        return [s for s in data if len(s.split()) >= min_word_count]
    elif all(isinstance(i, list) for i in data):
        # If data is a list of lists of strings, filter out strings in each list that are under the specified word count
        return [[s for s in sublist if len(s.split()) >= min_word_count] for sublist in data]
    else:
        raise ValueError("Input data should be a list of strings or a list of lists of strings.")
    
# Function to split a string into chunks
def split_string(s, chunk_size):
    """Split a string into chunks of a specified word count."""
    words = s.split()
    return [' '.join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size)]

# Function to split patient speech into even chunks
def split_into_chunks(data, chunk_size=100, n_chunks=10):
    """Split strings into chunks of a specified word count.
    Args:
        data (list): A list of strings or a list of lists of strings.
        chunk_size (int): The size of each chunk. (mutually exclusive with n_chunks)
        n_chunks (int): The number of chunks to split each string into. (mutually exclusive with chunk_size)
    Returns:
        A list of strings or a list of lists of strings, where the strings within each list have been combined and then split every chunk_size words or into n_chunks chunks."""

    if chunk_size:
        if all(isinstance(i, str) for i in data):
            # If data is a list of strings, combine the strings and split into chunks
            combined_string = ' '.join(data)
            return split_string(combined_string, chunk_size)
        elif all(isinstance(i, list) for i in data):
            # If data is a list of lists of strings, combine the strings in each list and split into chunks
            return [split_string(' '.join(sublist), chunk_size) for sublist in data]
        else:
            raise ValueError("Input data should be a list of strings or a list of lists of strings.")
    
    elif n_chunks:
        if all(isinstance(i, str) for i in data):
            # If data is a list of strings, combine all strings and split the result into chunks
            combined_string = ' '.join(data)
            return split_string(combined_string, len(combined_string.split()) // n_chunks)
        elif all(isinstance(i, list) for i in data):
            # If data is a list of lists of strings, combine all strings in each list and split the result into chunks
            return [split_string(' '.join(sublist), len(' '.join(sublist).split()) // n_chunks) for sublist in data]
        else:
            raise ValueError("Input data should be a list of strings or a list of lists of strings.")
    else:
        raise ValueError("You must specify either chunk_size or n_chunks, but not both.")

# Function to load patient and therapist speech turns from all documents in a folder and split into chunks  of turns with the minimum of a specified word count
def load_and_chunk_speech_turns(folder_path, min_word_count=250):
    """Load speech turns from all documents in a folder and combine them into chunks of a specified minimum word count.
    Args:
        folder_path (str): The path to the folder containing the documents.
        min_word_count (int): The minimum word count of each chunk.
    Returns:
        A list of lists of strings, where each list represents a document and each string within a list represents a chunk of speech turns."""

    # Initialize an empty list to hold all the chunks
    all_chunks = []

    # Compile the regular expression for matching the prefixes
    prefix_re = re.compile(r'^(P\d*:|P:|PATIENT:|P;|PATIENT;|T\d*:|T:|THERAPIST:|T;|THERAPIST;)')

    # Iterate over all files in the folder
    for filename in os.listdir(folder_path):
        # Check if the file is a Word document
        if filename.endswith('.docx'):
            # Load the document
            doc = Document(os.path.join(folder_path, filename))

            # Extract the text of each paragraph
            paragraphs = [p.text for p in doc.paragraphs]

            # Filter the paragraphs to only include those that start with the prefix
            paragraphs = [p for p in paragraphs if prefix_re.match(p)]

            # Combine the paragraphs into chunks of the specified minimum word count
            chunks = []
            chunk = []
            word_count = 0
            for paragraph in paragraphs:
                paragraph_word_count = len(paragraph.split())
                if word_count + paragraph_word_count >= min_word_count:
                    chunk.append(paragraph)
                    chunks.append(' '.join(chunk))
                    chunk = []
                    word_count = 0
                else:
                    chunk.append(paragraph)
                    word_count += paragraph_word_count
            if chunk:
                chunks.append(' '.join(chunk))

            # Add the chunks to the list of all chunks
            all_chunks.append(chunks)

    return all_chunks