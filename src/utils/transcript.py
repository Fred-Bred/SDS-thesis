"""
Utility functions for working with transcripts.
"""

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pandas as pd
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer

import xml.etree.ElementTree as ET
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from lxml import etree
import zipfile

import pandas as pd

import random
from faker import Faker

# Function to extract text from a docx file
def get_docx_text(filename, class_df=False):
    """Extract text from a docx file"""
    document = Document(filename)

    docText = '\n\n'.join([paragraph.text for paragraph in document.paragraphs])

    if class_df:
        doc_class = filename.split('_')[-1].split('.')[0]
        return pd.DataFrame({'text': [docText], 'class': [doc_class]})
    else:
        return docText

# Function to extract comments from a docx file
def extract_comments(filename):
    """Extract all comments from a docx file"""
    doc = Document(filename)
    try:
        comments_part = doc.part.package.part_related_by('http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
    except KeyError:
        return []
    root = ET.fromstring(comments_part.blob)
    comments = [elem.text for elem in root.iter() if elem.tag.endswith('comment')]
    return comments

# Function to extract all the comments of document 
def get_document_comments(docxFileName):
    """Returns a dictionary with comment id as key and comment string as value"""
    ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    comments_dict={}
    docxZip = zipfile.ZipFile(docxFileName)
    commentsXML = docxZip.read('word/comments.xml')
    et = etree.XML(commentsXML)
    comments = et.xpath('//w:comment',namespaces=ooXMLns)
    for c in comments:
        comment=c.xpath('string(.)',namespaces=ooXMLns)
        comment_id=c.xpath('@w:id',namespaces=ooXMLns)[0]
        comments_dict[comment_id]=comment
    return comments_dict

#Function to fetch all the comments in a paragraph
def paragraph_comments(paragraph, comments_dict):
    """Returns a list of comments in a given paragraph"""
    ooXMLns = {'w':'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    comments=[]
    for run in paragraph.runs:
        comment_reference=run._r.xpath("./w:commentReference")
        if comment_reference:
            comment_id=comment_reference[0].xpath('@w:id',namespaces=ooXMLns)[0]
            comment=comments_dict[comment_id]
            comments.append(comment)
    return comments

# Function to fetch all paragraphs with comments from a document
def comments_with_reference_paragraph(docxFileName):
    """Returns a dict with keys as paragraphs and values as comments"""
    document = Document(docxFileName)
    comments_dict=get_document_comments(docxFileName)
    comments_with_their_reference_paragraph={}
    for paragraph in document.paragraphs:  
        if comments_dict: 
            comments=paragraph_comments(paragraph,comments_dict)  
            if comments:
                comments_with_their_reference_paragraph[paragraph.text] = comments
    return comments_with_their_reference_paragraph

# Function to extract lists of text and comments from a dict where keys are paragraphs and values are comments
def extract_text_and_comments(comments_with_reference_paragraph, as_lists=False):
    """Returns a list or df of text and comments from a dict where keys are paragraphs and values are comments"""
    text_list = []
    comments_list = []
    for item in comments_with_reference_paragraph:
        for key, value in item.items():
            text_list.append(key)
            comments_list.append(value)
    # text_list = [item for sublist in text_list for item in sublist]
    comments_list = [item for sublist in comments_list for item in sublist]
    # Split label strings by comma
    comments_list = [i.split(", ") for i in comments_list]
    if as_lists:
        return text_list, comments_list
    else:
        # Create df
        df = pd.DataFrame(list(zip(text_list, comments_list)), columns = ["sentences", "labels"])
        
        # Convert the list of labels into a series of strings (so that get_dummies can process it)
        df['labels'] = df['labels'].apply(lambda x: ','.join(map(str, x)))

        # Convert the labels column into dummy variables
        df = df.join(df['labels'].str.get_dummies(','))

        # Drop the labels column
        df = df.drop(columns=['labels'])

        return df


# Function to create fake data
# Function to create a random Word document
def create_random_word_document(filename, n_docs=1):
    """Create n random Word documents"""
    fake = Faker()

    for i in range(n_docs):
        doc = Document()

        # Add 3-6 paragraphs
        for _ in range(random.randint(3, 6)):
            paragraph = doc.add_paragraph(fake.text())
            paragraph.add_comment(str(random.randint(0, 59)), author=fake.name(), initials='ab')

        doc.save(filename)